"""PRD.md → PRD.docx 변환 스크립트"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent.parent
SRC  = ROOT / "PRD.md"
DEST = ROOT / "PRD.docx"

# ── 브랜드 색상 ──────────────────────────────────────────────────
FOREST      = RGBColor(0x2D, 0x5A, 0x27)
FOREST_DARK = RGBColor(0x1A, 0x38, 0x16)
LIGHT_BG    = RGBColor(0xF0, 0xF5, 0xEF)
GRAY        = RGBColor(0x55, 0x65, 0x54)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BORDER_CLR  = "2D5A27"

def hex_color(element, hex_str):
    """XML 요소에 색상 설정"""
    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), hex_str)
    element.append(clr)

def set_cell_bg(cell, hex_str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_str)
    tcPr.append(shd)

def set_table_border(table):
    tbl    = table._tbl
    tblPr  = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "C8D8C5")
        borders.append(el)
    tblPr.append(borders)

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size   = Pt(20)
        run.font.color.rgb = FOREST_DARK
        # 하단 경계선
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "4")
        bot.set(qn("w:color"), BORDER_CLR)
        pBdr.append(bot)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size   = Pt(14)
        run.font.color.rgb = FOREST
    elif level == 3:
        run.font.size   = Pt(12)
        run.font.color.rgb = GRAY
        run.italic = True
    return p

def add_body(doc, text):
    # 인라인 **bold** 파싱
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
            r.font.size = Pt(10.5)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name  = "Courier New"
            r.font.size  = Pt(9.5)
            r.font.color.rgb = FOREST
        else:
            r = p.add_run(part)
            r.font.size = Pt(10.5)
    return p

def add_bullet(doc, text, level=0):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.6)
    p.paragraph_format.space_after  = Pt(3)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
            r.font.size = Pt(10.5)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name  = "Courier New"
            r.font.size  = Pt(9.5)
            r.font.color.rgb = FOREST
        else:
            r = p.add_run(part)
            r.font.size = Pt(10.5)
    return p

def parse_inline(text: str) -> list[tuple[str, dict]]:
    """**bold**, `code` 파싱 → [(text, {bold, code})]"""
    result = []
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if part.startswith("**") and part.endswith("**"):
            result.append((part[2:-2], {"bold": True}))
        elif part.startswith("`") and part.endswith("`"):
            result.append((part[1:-1], {"code": True}))
        else:
            result.append((part, {}))
    return result

def add_row_to_table(table, cells_text, is_header=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.paragraphs[0].clear()
        p    = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        if is_header:
            set_cell_bg(cell, "2D5A27")
            run = p.add_run(text.strip())
            run.bold           = True
            run.font.color.rgb = WHITE
            run.font.size      = Pt(10)
        else:
            parts = parse_inline(text.strip())
            for txt, fmt in parts:
                run = p.add_run(txt)
                run.font.size = Pt(10)
                if fmt.get("bold"):
                    run.bold = True
                if fmt.get("code"):
                    run.font.name  = "Courier New"
                    run.font.size  = Pt(9)
                    run.font.color.rgb = FOREST
    return row

def build_docx(src: Path, dest: Path):
    doc = Document()

    # ── 페이지 여백 ────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    # ── 기본 스타일 ────────────────────────────────────────────────
    nml = doc.styles["Normal"]
    nml.font.name = "맑은 고딕"
    nml.font.size = Pt(10.5)

    lines = src.read_text(encoding="utf-8").splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # 빈 줄
        if not line.strip():
            i += 1
            continue

        # 수평선
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "4")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "C8D8C5")
            pBdr.append(bot)
            pPr.append(pBdr)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # 제목
        if line.startswith("# "):
            # 문서 제목 — 커버 스타일
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after  = Pt(8)
            run = p.add_run(line[2:].strip())
            run.bold           = True
            run.font.size      = Pt(24)
            run.font.color.rgb = FOREST_DARK
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 1)
            i += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
            i += 1
            continue
        if line.startswith("#### "):
            add_heading(doc, line[5:].strip(), 3)
            i += 1
            continue

        # 테이블
        if line.startswith("|"):
            # 테이블 행 수집
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # 구분선 제거
            table_lines = [l for l in table_lines if not re.match(r"^\|[-| :]+\|$", l)]
            if not table_lines:
                continue
            # 컬럼 파싱
            def parse_row(l):
                return [c.strip() for c in l.strip("|").split("|")]
            rows = [parse_row(l) for l in table_lines]
            ncols = max(len(r) for r in rows)
            # 컬럼 너비 배분
            available = 14.0  # cm (A4 - margins)
            col_width  = available / ncols
            table = doc.add_table(rows=0, cols=ncols)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            set_table_border(table)
            # 컬럼 너비
            for col in table.columns:
                for cell in col.cells:
                    cell.width = Inches(col_width / 2.54)
            for j, row in enumerate(rows):
                # 부족한 셀 채우기
                while len(row) < ncols:
                    row.append("")
                add_row_to_table(table, row, is_header=(j == 0))
            doc.add_paragraph()
            continue

        # 순번 목록 (1. 2. 3.)
        m = re.match(r"^(\d+)\.\s+(.*)", line)
        if m:
            p   = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(3)
            for txt, fmt in parse_inline(m.group(2)):
                run = p.add_run(txt)
                run.font.size = Pt(10.5)
                if fmt.get("bold"):  run.bold = True
                if fmt.get("code"):
                    run.font.name = "Courier New"; run.font.size = Pt(9.5)
                    run.font.color.rgb = FOREST
            i += 1
            continue

        # 불릿 목록 (- * 들여쓰기 체크)
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            add_bullet(doc, m.group(2), indent)
            i += 1
            continue

        # 이탤릭 한 줄 (문서 끝 주석)
        if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            p   = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.italic         = True
            run.font.size      = Pt(9)
            run.font.color.rgb = GRAY
            p.paragraph_format.space_before = Pt(12)
            i += 1
            continue

        # 일반 본문
        add_body(doc, line.strip())
        i += 1

    doc.save(str(dest))
    print(f"OK: {dest}")

if __name__ == "__main__":
    build_docx(SRC, DEST)
